#!/usr/bin/env python3
"""
Génère les rapports journaliers FLOTTE au format DOCX (lundi 2 au lundi 9 février 2025).
Équipe : KOUAKOU EBOUHO (Scrum Master), KONAN OCEANE (Product Owner),
         YOBOUKOI DIVINE, DJEDJE ESTHER.
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'rapports_journaliers')
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Couleurs (thème FLOTTE — beige / marron professionnel)
COULEUR_TITRE = RGBColor(0x2d, 0x5a, 0x4a)      # vert bleuté foncé
COULEUR_SECTION = RGBColor(0x6f, 0x5d, 0x45)    # marron
COULEUR_SOUS_TEXTE = RGBColor(0x53, 0x46, 0x36) # marron foncé
COULEUR_TEXTE = RGBColor(0x2d, 0x2d, 0x2d)      # noir doux

# Contenu structuré : pour chaque date, listes de chaînes pour puces
REPORTS = {
    date(2025, 2, 2): {
        "realisations": [
            "Réunion de démarrage de la semaine d'amélioration.",
            "Analyse des quatre axes prioritaires : interface utilisateur, charges d'importation, pièces importées, location avec contraventions.",
            "KOUAKOU EBOUHO (Scrum Master) a animé la réunion et réparti les rôles.",
            "KONAN OCEANE (Product Owner) a formalisé le backlog et les critères d'acceptation.",
            "YOBOUKOI DIVINE et DJEDJE ESTHER ont analysé le code existant (modèles Vehicule, Location, ImportDemarche, templates, styles) et rédigé les spécifications détaillées.",
        ],
        "repartition": [
            "KOUAKOU EBOUHO : animation quotidienne, suivi du tableau de bord, déblocage.",
            "KONAN OCEANE : priorisation et validation des livrables.",
            "YOBOUKOI DIVINE : évolutions UI (sidebar, couleurs) et participation charges d'import.",
            "DJEDJE ESTHER : spécifications pièces importées et location (frais, contraventions), puis implémentation.",
        ],
        "ameliorations": [
            "État actuel : l'application gère déjà le parc véhicules, les démarches d'import (sans coûts détaillés), les locations (loyer, km, CT, assurance) mais sans coût global ni contraventions.",
            "Les couleurs véhicule sont en champs texte libres ; la sidebar n'est pas sticky.",
            "Aucune erreur bloquante signalée.",
        ],
        "prevision": [
            "Démarrer l'implémentation UI : sidebar sticky (CSS), revue des couleurs, préparation du paramétrage des couleurs véhicule en liste déroulante.",
        ],
    },
    date(2025, 2, 3): {
        "realisations": [
            "Mise en place de la barre latérale (sidebar) sticky : la sidebar reste fixe lors du défilement (position: sticky dans styles.css).",
            "Tests sur les pages tableau de bord, parc et chiffre d'affaires.",
            "Revue de la palette de couleurs (variables CSS) pour améliorer le contraste et l'accessibilité.",
            "KONAN OCEANE a validé les maquettes couleurs.",
            "YOBOUKOI DIVINE a implémenté le CSS sidebar et les variables ; DJEDJE ESTHER a documenté les changements et préparé la liste des couleurs pour le select véhicule.",
        ],
        "repartition": [
            "KOUAKOU EBOUHO : revue de code et vérification du comportement sticky sur plusieurs écrans.",
            "KONAN OCEANE : validation des couleurs et critères pour le select couleur.",
            "YOBOUKOI DIVINE : finalisation CSS, correction du dépassement sur mobile.",
            "DJEDJE ESTHER : liste des couleurs (Blanc, Noir, Gris, Bleu, Rouge, Argent, Vert, Beige, etc.) pour le paramétrage.",
        ],
        "ameliorations": [
            "Correction : sur une résolution intermédiaire, la sidebar provoquait un scroll horizontal ; ajout de min-width: 0 sur .main et overflow sur .content.",
            "Les couleurs sont désormais plus lisibles sur fond clair.",
        ],
        "prevision": [
            "Implémenter la couleur de voiture en select : paramétrage ou liste fixe, mise à jour du formulaire véhicule et des vues associées.",
        ],
    },
    date(2025, 2, 4): {
        "realisations": [
            "Mise en place du champ couleur en liste déroulante pour les véhicules.",
            "Choix retenu : paramétrage (table CouleurVehicule ou liste fixe) ; les champs couleur extérieure et intérieure utilisent désormais un Select avec options prédéfinies.",
            "KOUAKOU EBOUHO a ajouté le paramétrage des couleurs et modifié VehiculeForm.",
            "YOBOUKOI DIVINE a mis à jour le template vehicule_form.html et les vues ; KONAN OCEANE a validé la liste des couleurs et l'ergonomie.",
        ],
        "repartition": [
            "KOUAKOU EBOUHO : modèle/paramétrage couleurs, formulaire Django.",
            "KONAN OCEANE : recette et validation.",
            "YOBOUKOI DIVINE : templates et tests manuels.",
            "DJEDJE ESTHER : mise à jour des écrans détail véhicule pour un affichage cohérent.",
        ],
        "ameliorations": [
            "Aucune régression sur les véhicules existants : les anciennes valeurs en texte libre restent affichées ; les nouveaux enregistrements utilisent le select.",
            "Possibilité d'ajouter une entrée « Autre » avec champ texte optionnel ultérieurement.",
        ],
        "prevision": [
            "Attaquer le premier axe métier : charges d'importation (fret, dédouanement, frais transitaire, coût total). Conception du modèle et des écrans.",
        ],
    },
    date(2025, 2, 5): {
        "realisations": [
            "Conception et implémentation de la gestion détaillée des coûts d'importation.",
            "Nouveau modèle lié au véhicule avec : fret (transport international), frais de dédouanement, frais de service du transitaire, coût total (calcul automatique ou saisie manuelle).",
            "Formulaire et vues liste/détail créés.",
            "DJEDJE ESTHER et YOBOUKOI DIVINE ont développé le modèle, la migration, les vues et le formulaire ; KONAN OCEANE a rédigé les libellés et validé les calculs ; KOUAKOU EBOUHO a supervisé l'intégration et les permissions.",
        ],
        "repartition": [
            "KOUAKOU EBOUHO : permissions, URLs, intégration au menu Import.",
            "KONAN OCEANE : rédaction des labels et validation.",
            "YOBOUKOI DIVINE : modèles et formulaires.",
            "DJEDJE ESTHER : vues, templates et affichage sur la fiche véhicule.",
        ],
        "ameliorations": [
            "Correction : arrondi des montants en FCFA (entiers) pour éviter les décimales sur le coût total.",
            "Le coût total peut être recalculé à la volée ou figé selon besoin métier.",
        ],
        "prevision": [
            "Enchaîner sur les pièces ou parties de véhicules importées : modèle, quantités, coûts et association aux véhicules.",
        ],
    },
    date(2025, 2, 6): {
        "realisations": [
            "Mise en place de la gestion des pièces ou parties de véhicules importées.",
            "Nouveau modèle avec : désignation, quantité, coût unitaire, coût total, lien optionnel vers un véhicule (ForeignKey).",
            "Écrans de liste et de création/édition, avec filtre par véhicule ; association des pièces aux véhicules concernés.",
            "KOUAKOU EBOUHO et KONAN OCEANE ont défini le schéma et les règles métier ; YOBOUKOI DIVINE a implémenté le modèle et la migration ; DJEDJE ESTHER a créé les vues, formulaires, templates et le lien depuis la fiche véhicule.",
        ],
        "repartition": [
            "KOUAKOU EBOUHO : modèles et relations, admin Django.",
            "KONAN OCEANE : spécifications et recette.",
            "YOBOUKOI DIVINE : backend (modèle, formulaire).",
            "DJEDJE ESTHER : vues, templates, liste et détail véhicule.",
        ],
        "ameliorations": [
            "Une pièce peut être associée à un véhicule ou rester « orpheline » (véhicule null) pour les stocks généraux.",
            "Affichage du coût total (quantité × coût unitaire) automatique.",
        ],
        "prevision": [
            "Démarrer l'axe location : coût de location, frais annexes (assurance, entretien, carburant, etc.) et préparation de la gestion des contraventions.",
        ],
    },
    date(2025, 2, 7): {
        "realisations": [
            "Évolution du module location : ajout des champs coût de location et frais annexes (assurance, entretien, carburant, etc.) sur le modèle Location.",
            "Formulaire de location mis à jour.",
            "Début de la conception des contraventions : modèle avec montant, date, lieu, lien à la location, pénalité éventuelle.",
            "KONAN OCEANE a précisé les règles métier ; YOBOUKOI DIVINE et DJEDJE ESTHER ont implémenté les champs location et le squelette du modèle contravention ; KOUAKOU EBOUHO a animé le point technique et vérifié la cohérence avec le CA.",
        ],
        "repartition": [
            "KOUAKOU EBOUHO : coordination et cohérence CA.",
            "KONAN OCEANE : règles métier et critères d'acceptation.",
            "YOBOUKOI DIVINE : modèle Location étendu, formulaire.",
            "DJEDJE ESTHER : modèle Contravention, liaison à Location.",
        ],
        "ameliorations": [
            "Les frais annexes sont saisis en FCFA.",
            "Le coût total de la location (loyer + frais + pénalités) sera finalisé le lendemain.",
        ],
        "prevision": [
            "Finaliser contraventions : enregistrement des amendes, pénalités sur la facture de location, calcul du coût total et affichage sur la fiche location.",
        ],
    },
    date(2025, 2, 8): {
        "realisations": [
            "Finalisation de la gestion des contraventions : enregistrement des amendes (montant, date, référence, lien à la location).",
            "Ajout des pénalités liées à la facture de location ; calcul du coût total : loyer + frais annexes + pénalités/contraventions.",
            "Affichage du coût total sur la fiche location et dans la liste ; formulaires et vues de saisie des contraventions depuis la fiche location.",
            "Toute l'équipe a participé : KOUAKOU EBOUHO (intégration, tests), KONAN OCEANE (recette), YOBOUKOI DIVINE et DJEDJE ESTHER (développement des vues et calculs).",
        ],
        "repartition": [
            "KOUAKOU EBOUHO : calcul coût total, tests de bout en bout.",
            "KONAN OCEANE : validation fonctionnelle.",
            "YOBOUKOI DIVINE : formulaire contravention, affichage.",
            "DJEDJE ESTHER : agrégation des montants et mise à jour du détail location.",
        ],
        "ameliorations": [
            "Correction : prise en compte des locations sans frais annexes ni contraventions (valeurs nulles = 0).",
            "Le coût total est recalculé à chaque affichage à partir des données liées.",
        ],
        "prevision": [
            "Recette globale, corrections mineures éventuelles, synthèse de la semaine et rapport final. Préparation de la livraison des quatre axes.",
        ],
    },
    date(2025, 2, 9): {
        "realisations": [
            "Recette globale des quatre axes d'amélioration.",
            "L'équipe a ajouté des fonctionnalités et amélioré l'application tout au long de la semaine.",
            "Nouvelles fonctionnalités : interface (sidebar sticky, revue des couleurs, couleur véhicule en select), charges d'importation (fret, dédouanement, frais transitaire, coût total), pièces importées (quantités, coûts, association aux véhicules), location (coût, frais annexes, contraventions, pénalités, impact sur le coût total).",
            "Améliorations : palette de couleurs plus lisible, corrections d'affichage, libellés et messages de succès.",
            "KOUAKOU EBOUHO a piloté la recette et documenté les évolutions ; KONAN OCEANE a validé le backlog et clos les user stories ; YOBOUKOI DIVINE et DJEDJE ESTHER ont appliqué les correctifs et mis à jour la documentation.",
        ],
        "repartition": [
            "KOUAKOU EBOUHO : synthèse, rapport final, mise à jour README/documentation.",
            "KONAN OCEANE : validation produit et clôture du sprint.",
            "YOBOUKOI DIVINE et DJEDJE ESTHER : correctifs et tests de régression.",
        ],
        "ameliorations": [
            "Au cours de la semaine : ajout de fonctionnalités (sidebar sticky, couleur en select, charges d'import détaillées, pièces importées, location avec frais et contraventions) et améliorations (revue des couleurs, corrections d'erreurs, ergonomie).",
            "Tous les objectifs sont atteints. Aucun bug bloquant.",
            "Proposition pour la suite : exports PDF/Excel des coûts d'import et des locations.",
        ],
        "prevision": [
            "Livraison des rapports journaliers et de la version améliorée de l'application. Planification de la prochaine itération si besoin.",
        ],
    },
}

# Rapport combiné final (8–9 fév. + toutes les modifs) — 2 pages, dernier rapport journalier
SYNTHESE_FINALE = {
    "realisations": [
        "Rapport combiné des journées des 8 et 9 février : recette des quatre axes (interface, charges d'import, pièces importées, location avec contraventions), synthèse et livraison.",
        "Évolutions complémentaires intégrées : API REST (rate limiting, CORS), audit étendu (signals sur Vente, Dépense, Facture, etc. ; consultation et export par les admins).",
        "Champs fichier sur Facture, DocumentVehicule et ImportDemarche pour attacher les vrais documents (PDF, images).",
        "Fiche véhicule : totaux (charges d'import, dépenses, réparations), référence de coût pour la facture, mise en forme et UI plus professionnelle.",
        "Pays d'origine : liste complète des pays (datalist) avec saisie au fil de la frappe ; validation pour éviter la saisie d'un montant par erreur.",
        "Formulaires : menus déroulants avec saisie libre (datalist) sur de nombreux champs — locataire, étapes/statuts import, garantie, état livraison, phase dépense, type document, type réparation, prestataire, type facture, fournisseur, lieu carburant — pour faciliter la saisie tout en permettant une valeur libre.",
    ],
    "repartition": [
        "KOUAKOU EBOUHO : pilotage, synthèse et cohérence des livrables.",
        "KONAN OCEANE : validation produit et clôture du sprint.",
        "YOBOUKOI DIVINE : évolutions interface, formulaires et correctifs.",
        "DJEDJE ESTHER : développements métier, audit, API et documentation.",
    ],
    "ameliorations": [
        "Location : coût total (loyer + frais annexes + contraventions) affiché et calculé correctement ; pas de régression sur les valeurs nulles.",
        "UX : champs avec suggestions (datalist) conservent la possibilité de saisie manuelle pour ne pas bloquer les cas particuliers.",
        "Documentation (ANALYSE_AVANT_AMELIORATIONS, NORMES, API) et rapports journaliers à jour.",
    ],
    "prevision": [
        "Livraison de l'application et des rapports. Planification de la prochaine itération si besoin.",
    ],
}


def _style_run(run, size=11, bold=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _style_run(run, size=18, bold=True, color=COULEUR_TITRE)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _style_run(run, size=11, color=COULEUR_SOUS_TEXTE)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_team_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _style_run(run, size=10, color=COULEUR_SOUS_TEXTE)
    p.paragraph_format.space_after = Pt(18)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _style_run(run, size=13, bold=True, color=COULEUR_SECTION)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0)
    return p


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        _style_run(run, size=11, color=COULEUR_TEXTE)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
    return doc


def build_report(doc_date):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    doc.add_paragraph()
    add_title(doc, "Projet FLOTTE — Rapport journalier")
    add_subtitle(doc, doc_date.strftime("%A %d %B %Y"))
    add_team_line(
        doc,
        "Équipe : KOUAKOU EBOUHO (Scrum Master) · KONAN OCEANE (Product Owner) · YOBOUKOI DIVINE · DJEDJE ESTHER"
    )

    data = REPORTS.get(doc_date)
    if not data:
        return doc

    add_section_heading(doc, "Réalisations du jour")
    add_bullet_list(doc, data["realisations"])

    add_section_heading(doc, "Répartition des tâches")
    add_bullet_list(doc, data["repartition"])

    section_ameliorations = "Améliorations et remarques"
    if doc_date == date(2025, 2, 9):
        section_ameliorations = "Fonctionnalités ajoutées et améliorations"
    add_section_heading(doc, section_ameliorations)
    add_bullet_list(doc, data["ameliorations"])

    add_section_heading(doc, "Prévision du lendemain" if doc_date != date(2025, 2, 9) else "Prévision suite")
    add_bullet_list(doc, data["prevision"])

    return doc


def build_synthese_finale():
    """Génère le rapport journalier combiné final (2 pages) — 8–9 fév. + toutes les modifs."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    doc.add_paragraph()
    add_title(doc, "Projet FLOTTE — Rapport journalier (synthèse finale)")
    add_subtitle(doc, "8–9 février 2026 — Dernier rapport journalier")
    add_team_line(
        doc,
        "Équipe : KOUAKOU EBOUHO (Scrum Master) · KONAN OCEANE (Product Owner) · YOBOUKOI DIVINE · DJEDJE ESTHER"
    )

    add_section_heading(doc, "Réalisations")
    add_bullet_list(doc, SYNTHESE_FINALE["realisations"])
    add_section_heading(doc, "Répartition des tâches")
    add_bullet_list(doc, SYNTHESE_FINALE["repartition"])
    add_section_heading(doc, "Améliorations et remarques")
    add_bullet_list(doc, SYNTHESE_FINALE["ameliorations"])
    add_section_heading(doc, "Prévision suite")
    add_bullet_list(doc, SYNTHESE_FINALE["prevision"])

    return doc


def main():
    for doc_date in sorted(REPORTS.keys()):
        doc = build_report(doc_date)
        name = f"Rapport_journalier_FLOTTE_{doc_date.isoformat()}.docx"
        path = os.path.join(OUTPUT_DIR, name)
        doc.save(path)
        print(f"Généré : {path}")

    # Rapport combiné final (8–9 fév. + toutes les modifs) — 2 pages
    doc_synthese = build_synthese_finale()
    path_synthese = os.path.join(OUTPUT_DIR, "Rapport_journalier_FLOTTE_synthese_finale.docx")
    doc_synthese.save(path_synthese)
    print(f"Généré : {path_synthese}")

    print(f"\n{len(REPORTS) + 1} rapport(s) dont 1 synthèse finale dans : {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
